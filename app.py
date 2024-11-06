from fastapi import FastAPI, UploadFile, File, HTTPException,Query
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings
from concurrent.futures import ThreadPoolExecutor
from typing import List
import aiofiles
import uuid
# Initialize FastAPI app
app = FastAPI()

# Load embedding model (CPU)
embedder = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')

# Initialize ChromaDB client with persistence settings
client = chromadb.Client(Settings(persist_directory="./chroma_db"))
collection = client.get_or_create_collection("documents")

# Define a document schema
# class Document(BaseModel):
#     id: str
#     content: str

# Executor for concurrency handling
executor = ThreadPoolExecutor(max_workers=10)

import os
from PyPDF2 import PdfReader
from docx import Document
from fastapi import HTTPException
from io import BytesIO

async def process_file(file: UploadFile) -> str:
    try:
        # Read the file as bytes
        content = await file.read()

        # Check for file type and extract text accordingly
        file_extension = os.path.splitext(file.filename)[1].lower()

        # Handle PDF files
        if file_extension == '.pdf':
            text = extract_text_from_pdf(content)

        # Handle DOCX files
        elif file_extension == '.docx':
            text = extract_text_from_docx(content)

        # Handle plain text files
        else:
            text = content.decode('utf-8', errors='ignore')  # Use 'ignore' to skip invalid characters

        if not text:
            raise HTTPException(status_code=400, detail=f"File {file.filename} has no text content.")

        return text
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing file {file.filename}: {str(e)}")

def extract_text_from_pdf(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content))
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error reading PDF: {str(e)}")

def extract_text_from_docx(content: bytes) -> str:
    try:
        doc = Document(BytesIO(content))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error reading DOCX: {str(e)}")

def generate_embeddings(text: str):
    try:
        # Use the SentenceTransformer model to generate embeddings
        embeddings = embedder.encode([text])  # Generate embeddings for the input text
        return embeddings[0].tolist()  # Return the first embedding (convert to list for JSON compatibility)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating embeddings: {str(e)}")

# @app.post("/ingest/")
# async def ingest_documents(files: List[UploadFile] = File(...)):
#     documents = []
#     for file in files:
#         print(f"Processing file: {file.filename}")
#         content = await process_file(file)
#         # if content:
#         #     documents.append(Document(id=file.filename, content=content))
#         #     print(f"Processed file: {file.filename}")
#         # else:
#         #     print(f"File {file.filename} is empty or could not be read")
       

#     # Ingest documents into ChromaDB
#     try:
#         embeddings = [embedder.encode(doc.content).tolist() for doc in documents]
#         for doc, embedding in zip(documents, embeddings):
#             collection.add(
#                 documents=[doc.content],
#                 metadatas=[{"id": doc.id}],
#                 embeddings=[embedding]
#             )
#         return JSONResponse(content={"message": "Documents ingested successfully"})
#     except Exception as e:
#         print(f"Error during ingestion: {e}")
#         raise HTTPException(status_code=500, detail="Error ingesting documents into ChromaDB")


@app.post("/ingest/")
async def ingest_documents(files: List[UploadFile] = File(...)):
    try:
        documents = []
        for file in files:
            # Process the file and get content
            content = await process_file(file)

            # Generate embeddings for the content
            embedding = generate_embeddings(content)  # Assume this is a function that generates embeddings

            document_id = str(uuid.uuid4())  # Generate a unique ID (UUID)

            # Add the document to ChromaDB collection
            collection.add(
                ids=[document_id],  # Unique identifier for each document
                documents=[content],  # The actual document text
                metadatas=[{"filename": file.filename}],  # Metadata (optional)
                embeddings=[embedding]  # Embedding for the document
             )
            documents.append(file.filename)

        return JSONResponse(content={"message": "Documents ingested successfully"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error ingesting documents: {str(e)}")
    

# @app.post("/query/")
# async def query_document(query: str = Query(..., description="Search query text")):
#     print(f"taking query_document {query}")
#     query_embedding = embedder.encode(query).tolist()
#     print("done query_document")
#     try:
#         results = collection.query(query_embeddings=[query_embedding], n_results=5)
#         return JSONResponse(content={"results": results})
#     except Exception as e:
#         raise HTTPException(status_code=500, detail="Error querying ChromaDB")


class QueryRequest(BaseModel):
    query: str

@app.post("/query/")
async def query_document(query_request: QueryRequest):
    query = query_request.query
    print(f"Processing query: {query}")
    query_embedding = embedder.encode(query).tolist()

    try:
        # Execute the query on the ChromaDB collection
        results = collection.query(query_embeddings=[query_embedding], n_results=5)

        # Return the results as JSON
        return {"results": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error querying ChromaDB: {str(e)}")



# Run FastAPI server
# Command: uvicorn app:app --reload
